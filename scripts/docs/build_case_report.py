from __future__ import annotations

# ruff: noqa: E501
import datetime as dt
import textwrap
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "output" / "doc" / "Prizma_полный_технический_отчет.docx"

ARTIFACT_GLOBS = [
    "README.md",
    "Makefile",
    "Dockerfile",
    "docker-compose*.yml",
    ".env.example",
    ".gitlab-ci.yml",
    ".github/workflows/*.yml",
    "pyproject.toml",
    "dvc.yaml",
    "params.yaml",
    "docs/**/*.md",
    "docs/api/openapi.json",
    "deploy/**/*.md",
    "deploy/**/*.yaml",
    "deploy/**/*.yml",
    "deploy/**/*.conf",
    "charts/prizma/**/*.yaml",
    "observability/**/*.yml",
    "observability/**/*.yaml",
    "observability/**/*.json",
    "scripts/**/*.py",
    "tests/**/*.py",
    "web/**/*",
    "triton-model-repository/**/*",
]

EXCLUDE_PARTS = {
    "__pycache__",
    ".DS_Store",
}


def current_git_commit() -> str:
    head = ROOT / ".git" / "HEAD"
    try:
        value = head.read_text(encoding="utf-8").strip()
        if value.startswith("ref: "):
            ref_path = ROOT / ".git" / value.removeprefix("ref: ")
            return ref_path.read_text(encoding="utf-8").strip()[:7]
        return value[:7]
    except OSError:
        return "недоступно"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(8.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, "D9EAD3")

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)


def add_code_block(document: Document, text: str, max_line: int = 132) -> None:
    if not text:
        text = "[пусто]"
    for raw_line in text.splitlines() or [""]:
        chunks = textwrap.wrap(
            raw_line,
            width=max_line,
            replace_whitespace=False,
            drop_whitespace=False,
        ) or [""]
        for line in chunks:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1
            run = paragraph.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(7)


def add_command(document: Document, command: str, expected: str) -> None:
    document.add_paragraph("Команда:", style="List Bullet")
    add_code_block(document, command)
    document.add_paragraph("Ожидаемый/полученный результат:", style="List Bullet")
    add_code_block(document, expected)


def artifact_paths() -> list[Path]:
    paths: set[Path] = set()
    for pattern in ARTIFACT_GLOBS:
        for path in ROOT.glob(pattern):
            if path.is_file() and not any(part in EXCLUDE_PARTS for part in path.parts):
                paths.add(path)
    return sorted(paths, key=lambda item: item.relative_to(ROOT).as_posix())


def read_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="utf-8", errors="replace")


def artifact_category(path: Path) -> str:
    rel = path.relative_to(ROOT).as_posix()
    if rel.startswith("docs/"):
        return "Документация и отчетные артефакты"
    if rel.startswith(".github/") or rel == ".gitlab-ci.yml":
        return "CI/CD"
    if rel.startswith("charts/") or rel.startswith("deploy/"):
        return "Kubernetes, Helm и deployment"
    if rel.startswith("observability/"):
        return "Observability"
    if rel.startswith("scripts/"):
        return "Automation/MLOps/Security scripts"
    if rel.startswith("tests/"):
        return "Тесты"
    if rel.startswith("web/"):
        return "Frontend"
    if rel.startswith("triton-model-repository/"):
        return "Triton"
    return "Корневой артефакт"


def parse_markdown_table(path: Path) -> tuple[list[str], list[list[str]]]:
    lines = [line.strip() for line in read_text(path).splitlines()]
    table_lines = [line for line in lines if line.startswith("|") and line.endswith("|")]
    if len(table_lines) < 2:
        return [], []

    headers = [item.strip() for item in table_lines[0].strip("|").split("|")]
    rows = []
    for line in table_lines[2:]:
        cells = [item.strip() for item in line.strip("|").split("|")]
        if len(cells) == len(headers):
            rows.append(cells)
    return headers, rows


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)
    for name, size, color in [
        ("Title", 24, "0B4F3D"),
        ("Heading 1", 18, "0B4F3D"),
        ("Heading 2", 14, "174A3A"),
        ("Heading 3", 11, "174A3A"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)


def add_title(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Prizma: полный технический отчет по кейсу")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(11, 79, 61)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Архитектура, MLOps, CI/CD, SRE-артефакты, отладка и проверки").italic = True

    meta_rows = [
        ["Дата генерации", dt.datetime.now().strftime("%Y-%m-%d %H:%M")],
        ["Репозиторий", "https://github.com/Pernya/prizma"],
        ["Локальный путь", str(ROOT)],
        ["Git commit на момент генерации", current_git_commit()],
        ["Основной домен", "prizma.pernyaev.ru"],
        ["API upstream", "api.pernyaev.ru через Cloudflare Tunnel"],
    ]
    add_table(document, ["Поле", "Значение"], meta_rows)


def add_summary(document: Document) -> None:
    document.add_heading("1. Краткое резюме", level=1)
    paragraphs = [
        "Проект Prizma доведен до production-oriented baseline: есть backend API, frontend, "
        "асинхронный worker, RabbitMQ, MinIO/S3-compatible storage, Triton-compatible контур, "
        "Prometheus/Grafana, Docker Compose, Helm/Kubernetes, GitHub Actions CI/CD и набор "
        "MLOps/SRE/security gates.",
        "Рабочий публичный режим: frontend развернут на VPS через Docker Compose и nginx, "
        "а backend может оставаться локально на Mac и публиковаться через Cloudflare Tunnel "
        "как api.pernyaev.ru. VPS проксирует /api/* на этот upstream.",
        "GitLab CI/CD подготовлен, но фактический CD был перенесен на GitHub Actions из-за "
        "блокировки GitLab runners верификацией аккаунта.",
    ]
    for text in paragraphs:
        document.add_paragraph(text)


def add_architecture(document: Document) -> None:
    document.add_heading("2. Архитектура и контуры", level=1)
    rows = [
        [
            "Frontend",
            "`web/`, nginx на VPS",
            "Загрузка изображения, выбор стиля, статус, результат, скачать/поделиться",
        ],
        [
            "Backend API",
            "FastAPI в `src/prizma_backend`",
            "Health/readiness, styles, jobs, result, metrics",
        ],
        ["Worker", "`prizma_backend.worker`", "Асинхронная обработка задач через RabbitMQ"],
        ["Queue", "RabbitMQ", "Буферизация тяжелых задач обработки"],
        ["Storage", "MinIO/S3-compatible", "Хранение input/output артефактов"],
        [
            "Inference",
            "local Pillow или Triton-compatible HTTP client",
            "MVP-обработка и готовность к GPU inference",
        ],
        ["Observability", "Prometheus + Grafana", "Метрики, dashboard, Prometheus alerts"],
        [
            "Kubernetes",
            "Helm chart + values",
            "Managed Kubernetes target, HPA, ServiceMonitor, ExternalSecret",
        ],
        ["CD", "GitHub Actions", "VPS frontend deploy и manual Kubernetes deploy"],
    ]
    add_table(document, ["Компонент", "Реализация", "Назначение"], rows)


def add_coverage(document: Document) -> None:
    document.add_heading("3. Матрица покрытия требований", level=1)
    headers, rows = parse_markdown_table(ROOT / "docs/coverage/requirements-traceability.md")
    if headers and rows:
        add_table(document, headers, rows)
    else:
        add_code_block(document, read_text(ROOT / "docs/coverage/requirements-traceability.md"))


def add_sre(document: Document) -> None:
    document.add_heading("4. SLI/SLO/SLA и бюджет ошибок", level=1)
    headers, rows = parse_markdown_table(ROOT / "docs/sre/slo.md")
    document.add_paragraph(
        "Ключевые SLI/SLO перенесены в отдельный русскоязычный SRE-документ. "
        "Для MVP публичный коммерческий SLA не заявлен; внутренний ориентир - 99,9% доступности API."
    )
    if headers and rows:
        add_table(document, headers, rows)
    document.add_paragraph("Полный текст SRE-документа включен в приложение с артефактами.")


def add_debug_log(document: Document) -> None:
    document.add_heading("5. Журнал проблем, решений и проверок", level=1)
    rows = [
        [
            "DNS и домен prizma.pernyaev.ru",
            "dig сначала не возвращал A-запись; пользователь не понимал, как без Cloudflare Tunnel попасть на VPS.",
            "DNS-запись была создана в Cloudflare, но нужно было дождаться распространения и проверить, что A-запись ведет на публичный IP VPS.",
            "Проверять `dig +short prizma.pernyaev.ru A`, затем `curl -I http://prizma.pernyaev.ru/`. Для обычного VPS Tunnel не нужен, если 80/443 доступны извне.",
        ],
        [
            "Публичный IP VPS",
            "Изначально использовался IP 95.163.244.138, затем в панели Reg.Cloud обнаружен фактический публичный IP 194.226.142.176.",
            "DNS/SSH/деплой должны ссылаться на актуальный публичный IP, а приватный 192.168.0.165 используется только внутри private network.",
            "Использовать `ssh root@194.226.142.176` и GitHub secret `VPS_HOST=194.226.142.176`.",
        ],
        [
            "SSH/SCP",
            "`Connection closed`, неверный запуск `scp` с удаленного сервера, проблемы с password/publickey.",
            "SCP надо запускать с Mac, потому что локальная папка находится на Mac. Для GitHub Actions нужен отдельный deploy key без passphrase.",
            "Добавлен public key в `~/.ssh/authorized_keys`, private key записан в GitHub secret `VPS_SSH_KEY`.",
        ],
        [
            "Docker Compose на VPS",
            "`docker-compose-plugin` не находился через apt; `docker -f` давал unknown shorthand flag; docker-compose v1 ломался на recreate с `ContainerConfig`.",
            "На VPS доступен legacy `docker-compose`. Workflow сделан portable: сначала пробует `docker compose`, затем `docker-compose`; перед up делает down и удаляет старые контейнеры.",
            "GitHub Actions VPS deploy стал успешным после правок workflow.",
        ],
        [
            "Cloudflare Tunnel / ingress",
            "Была путаница между Tunnel, ingress и DNS. Также обсуждалась настройка route на localhost:8000.",
            "Cloudflare Tunnel нужен для локального backend на Mac. VPS nginx проксирует `/api/*` на `https://api.pernyaev.ru`; Kubernetes Ingress - отдельный механизм для managed cluster.",
            "Проверять `https://api.pernyaev.ru/healthz` и route в Zero Trust Tunnel.",
        ],
        [
            "502 Bad Gateway",
            "При POST/загрузке пользователь видел HTML `502 Bad Gateway` от nginx.",
            "502 означает, что VPS nginx не получил корректный ответ от upstream `api.pernyaev.ru` или backend/tunnel временно недоступен. Были добавлены proxy timeouts и отключено request buffering для upload.",
            "Текущая проверка: `POST /api/v1/jobs` возвращает 202, status - 200, result - `image/png`.",
        ],
        [
            "Старый frontend после обновления",
            "После refresh пользователь не видел изменения frontend.",
            "Вероятная причина - cache HTML/JS. Добавлены `Cache-Control: no-cache` для HTML/JS/CSS/SVG/manifest на VPS nginx.",
            "Проверять заголовки `curl -I http://prizma.pernyaev.ru/app.js?v=20260603-2`.",
        ],
        [
            "GitLab",
            "GitLab push прошел, но runners/CD заблокированы верификацией телефона/карты.",
            "GitLab CI оставлен как подготовленный артефакт, рабочий контур CI/CD перенесен в GitHub Actions.",
            "Проверять GitHub Actions: CI, Deploy VPS Frontend.",
        ],
    ]
    add_table(document, ["Проблема", "Симптом", "Причина", "Решение и проверка"], rows)


def add_checks(document: Document) -> None:
    document.add_heading("6. Команды проверок и ожидаемый вывод", level=1)
    checks = [
        (
            "Проверка публичного frontend",
            "curl -sS -D - http://prizma.pernyaev.ru/ -o /tmp/prizma-index.html",
            "Ожидаемо: `HTTP/1.1 200 OK`, HTML содержит `<title>Prizma</title>`, manifest и `app.js?v=20260603-2`.",
        ),
        (
            "Проверка frontend bundle",
            "curl -sS -D - 'http://prizma.pernyaev.ru/app.js?v=20260603-2' -o /tmp/app.js",
            "Ожидаемо: `HTTP/1.1 200 OK`, `Content-Type: application/javascript`, `Cache-Control: no-cache` после деплоя cache-control фикса.",
        ),
        (
            "Проверка API upstream",
            "curl -sS -D - https://api.pernyaev.ru/healthz",
            'Ожидаемо: `HTTP/2 200`, тело `{"status":"ok"}`. Если timeout/502 - проверить backend на Mac и Cloudflare Tunnel.',
        ),
        (
            "Проверка API через VPS proxy",
            "curl -sS -D - http://prizma.pernyaev.ru/api/v1/styles",
            "Ожидаемо: `HTTP/1.1 200 OK`, JSON-массив стилей `noir`, `vivid`, `sketch`.",
        ),
        (
            "Smoke upload job",
            "curl -sS -D - -F style=noir -F file=@/tmp/prizma-smoke.png http://prizma.pernyaev.ru/api/v1/jobs",
            "Ожидаемо: `HTTP/1.1 202 Accepted`, JSON содержит `job_id`, `status=queued`, `status_url`.",
        ),
        (
            "Smoke status/result",
            "curl -sS -D - http://prizma.pernyaev.ru/api/v1/jobs/<job_id>\n"
            "curl -sS -D - -o /tmp/result.png http://prizma.pernyaev.ru/api/v1/jobs/<job_id>/result\n"
            "file /tmp/result.png",
            "Ожидаемо: status `succeeded`, result endpoint `HTTP/1.1 200 OK`, `Content-Type: image/png`, `file` показывает PNG.",
        ),
        (
            "Локальные Python проверки",
            "make lint\nmake test",
            "Ожидаемо: `All checks passed!`, `8 passed`.",
        ),
        (
            "Security gates",
            "make security",
            "Ожидаемо: `No known vulnerabilities found`, Bandit без findings, `No secrets detected`.",
        ),
        (
            "MLOps gates",
            "make mlops",
            "Ожидаемо: создаются `reports/mlops/data-validation.json`, `benchmark.json`, `drift.json`, обновляется `docs/mlops/model-card.md`.",
        ),
        (
            "OpenAPI contract drift",
            "make contract && git diff --exit-code docs/api/openapi.json",
            "Ожидаемо: команда завершается с exit code 0; если есть diff, schema изменилась и contract надо зафиксировать.",
        ),
        (
            "Kubernetes/Helm validation",
            "make kube-validate",
            "Ожидаемо: `Summary: 112 resources found in 12 files - Valid: 97, Invalid: 0, Errors: 0, Skipped: 15`.",
        ),
        (
            "GitHub Actions",
            "Открыть Actions в GitHub или проверить API GitHub Actions.",
            "Ожидаемо: workflow `CI` и `Deploy VPS Frontend` имеют статус `completed success` для актуального commit.",
        ),
    ]
    for title, command, expected in checks:
        document.add_heading(title, level=2)
        add_command(document, command, expected)


def add_artifact_registry(document: Document, paths: list[Path]) -> None:
    document.add_heading("7. Реестр артефактов", level=1)
    rows = []
    for path in paths:
        rel = path.relative_to(ROOT).as_posix()
        rows.append([rel, artifact_category(path)])
    add_table(document, ["Путь", "Категория"], rows)


def add_appendix(document: Document, paths: list[Path]) -> None:
    document.add_section(WD_SECTION_START.NEW_PAGE)
    document.add_heading("Приложение A. Полные тексты артефактов", level=1)
    document.add_paragraph(
        "Ниже приведены текстовые артефакты репозитория: документация, CI/CD, Helm/Kubernetes, "
        "observability, scripts, tests и frontend. Секреты не включались; secret examples являются шаблонами."
    )
    for path in paths:
        rel = path.relative_to(ROOT).as_posix()
        document.add_heading(rel, level=2)
        add_code_block(document, read_text(path))


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_document(document)
    add_title(document)
    document.add_page_break()
    add_summary(document)
    add_architecture(document)
    add_coverage(document)
    add_sre(document)
    add_debug_log(document)
    add_checks(document)
    paths = artifact_paths()
    add_artifact_registry(document, paths)
    add_appendix(document, paths)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
